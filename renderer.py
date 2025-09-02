from docx import Document

def generate_contract(template_path, context, clauses, styles, output_path):
    doc = Document(template_path)

    # Substitui placeholders simples do template
    for p in doc.paragraphs:
        for key, val in context.items():
            placeholder = "{{ " + str(key) + " }}"
            if placeholder in p.text:
                p.text = p.text.replace(placeholder, str(val))

    # Insere cláusulas no marcador {{ clausulas }}
    for p in list(doc.paragraphs):  # list() para iterar sobre snapshot
        if "{{ clausulas }}" in p.text:
            p.text = p.text.replace("{{ clausulas }}", "")
            for idx, c in enumerate(clauses, start=1):
                titulo = c.get("titulo", "")
                texto  = c.get("texto", "")
                # Usando format() em vez de f-string (evita erros com chaves)
                clause_text = "Cláusula {}ª — {}\n{}".format(idx, titulo, texto)
                para = doc.add_paragraph(clause_text)
                apply_style(para, styles.get("clausulas_extra", {}))

    doc.save(output_path)

def apply_style(paragraph, style_conf):
    """Aplica negrito/itálico/CAIXA ALTA e fonte em todos os runs do parágrafo."""
    bold   = bool(style_conf.get("bold"))
    italic = bool(style_conf.get("italic"))
    caps   = bool(style_conf.get("caps"))
    font   = style_conf.get("font", "Times New Roman")

    # Se for caixa alta, transforma todo o texto e recria o parágrafo
    if caps:
        full_text = paragraph.text.upper()
        for r in paragraph.runs:
            r.text = ""
        paragraph.add_run(full_text)

    for run in paragraph.runs:
        run.bold = bold
        run.italic = italic
        try:
            run.font.name = font
        except Exception:
            pass
